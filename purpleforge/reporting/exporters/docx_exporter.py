"""Word document (DOCX) report exporter using python-docx."""

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from purpleforge.utils.logging import get_logger

logger = get_logger(__name__)


def _has_docx() -> bool:
    """Check if python-docx is available."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def export_docx(
    report_data: Dict[str, Any],
    output_path: Path,
) -> Optional[Path]:
    """
    Export report as Word document (DOCX).

    Args:
        report_data: Report data dictionary
        output_path: Path for DOCX output

    Returns:
        Path to created DOCX or None if failed
    """
    if not _has_docx():
        logger.warning("python-docx not installed. Install with: pip install python-docx")
        return None

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # Set document properties
        core_props = doc.core_properties
        core_props.title = f"PurpleForge Report - {report_data.get('run_id', 'Unknown')}"
        core_props.author = "PurpleForge"
        core_props.subject = "Security Assessment Report"

        # ===== Title Page =====
        title = doc.add_heading("PurpleForge", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading("Security Assessment Report", 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Report metadata table
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = 'Table Grid'

        meta_data = [
            ("Run ID", report_data.get("run_id", "N/A")),
            ("Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")),
            ("Status", report_data.get("status", "completed")),
            ("Scenario", report_data.get("scenario_name", "N/A")),
            ("Tool Version", report_data.get("tool_version", "0.1.0")),
        ]

        for i, (label, value) in enumerate(meta_data):
            meta_table.rows[i].cells[0].text = label
            meta_table.rows[i].cells[1].text = str(value)
            # Bold the labels
            meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_page_break()

        # ===== Executive Summary =====
        doc.add_heading("Executive Summary", 1)

        summary_text = f"""This report documents a controlled security assessment scenario executed using PurpleForge.

Assessment Period: {report_data.get('started_at', 'N/A')} to {report_data.get('completed_at', 'N/A')}
Scenario: {report_data.get('scenario_name', 'N/A')}
Category: {report_data.get('category', 'web')}
Mode: {report_data.get('mode', 'assess')}
"""
        doc.add_paragraph(summary_text)

        # Coverage stats
        doc.add_heading("Coverage Statistics", 2)

        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'

        stats = [
            ("Total Steps", str(report_data.get("total_steps", 0))),
            ("Steps with Evidence", str(report_data.get("steps_with_evidence", 0))),
            ("Coverage", f"{report_data.get('coverage', 0)}%"),
            ("Timeline Events", str(len(report_data.get("timeline", [])))),
        ]

        for i, (label, value) in enumerate(stats):
            stats_table.rows[i].cells[0].text = label
            stats_table.rows[i].cells[1].text = value
            stats_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # ===== Target Information =====
        doc.add_heading("Target Information", 1)

        target_para = doc.add_paragraph()
        target_para.add_run("Target URL: ").bold = True
        target_para.add_run(report_data.get("target_url", "N/A"))

        verif_para = doc.add_paragraph()
        verif_para.add_run("Verification Status: ").bold = True
        verif_para.add_run(report_data.get("verification_status", "unknown"))

        # ===== Steps Executed =====
        doc.add_heading("Steps Executed", 1)

        steps = report_data.get("steps", [])
        if steps:
            steps_table = doc.add_table(rows=len(steps) + 1, cols=4)
            steps_table.style = 'Table Grid'

            # Header row
            headers = ["Step ID", "Type", "Description", "Evidence"]
            for j, header in enumerate(headers):
                steps_table.rows[0].cells[j].text = header
                steps_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

            # Data rows
            for i, step in enumerate(steps, 1):
                steps_table.rows[i].cells[0].text = step.get("id", "")
                steps_table.rows[i].cells[1].text = step.get("type", "")
                steps_table.rows[i].cells[2].text = step.get("description", "")[:50]
                evidence = "Yes" if step.get("has_evidence") else "No"
                steps_table.rows[i].cells[3].text = evidence

        doc.add_paragraph()

        # ===== ATT&CK Techniques =====
        attack_techniques = report_data.get("attack_techniques", [])
        if attack_techniques:
            doc.add_heading("ATT&CK Mapping", 1)

            for tech in attack_techniques:
                doc.add_heading(f"{tech.get('id', '')} - {tech.get('name', '')}", 2)

                tactic_para = doc.add_paragraph()
                tactic_para.add_run("Tactic: ").bold = True
                tactic_para.add_run(tech.get("tactic", ""))

                desc = tech.get("description", "")
                if len(desc) > 300:
                    desc = desc[:300] + "..."
                doc.add_paragraph(desc)

                mitigations = tech.get("mitigations", [])
                if mitigations:
                    doc.add_paragraph("Mitigations:", style='List Bullet')
                    for m in mitigations[:5]:
                        doc.add_paragraph(
                            f"{m.get('id', '')} - {m.get('name', '')}",
                            style='List Bullet'
                        )

                url_para = doc.add_paragraph()
                url_para.add_run("Reference: ").bold = True
                url_para.add_run(tech.get("url", ""))

                doc.add_paragraph()

        # ===== Detection Recommendations =====
        doc.add_heading("Detection Recommendations", 1)

        recommendations = report_data.get("detection_recommendations", "")
        if recommendations:
            # Simple conversion of markdown to paragraphs
            for line in recommendations.split("\n"):
                line = line.strip()
                if line.startswith("###"):
                    doc.add_heading(line.replace("#", "").strip(), 3)
                elif line.startswith("##"):
                    doc.add_heading(line.replace("#", "").strip(), 2)
                elif line.startswith("-"):
                    doc.add_paragraph(line[1:].strip(), style='List Bullet')
                elif line.startswith("1.") or line.startswith("2.") or line.startswith("3."):
                    doc.add_paragraph(line[2:].strip(), style='List Number')
                elif line:
                    doc.add_paragraph(line)
        else:
            doc.add_paragraph("See Markdown report for detailed detection recommendations.")

        # ===== Timeline Summary =====
        timeline = report_data.get("timeline", [])
        if timeline:
            doc.add_heading("Timeline Summary", 1)

            doc.add_paragraph(f"Total events: {len(timeline)}")

            # Show first 10 events
            for event in timeline[:10]:
                event_para = doc.add_paragraph()
                event_para.add_run(f"{event.get('timestamp', '')} - ").bold = True
                event_para.add_run(f"{event.get('type', '')} ")
                if event.get("step_id"):
                    event_para.add_run(f"(Step: {event['step_id']})")

            if len(timeline) > 10:
                doc.add_paragraph(f"... and {len(timeline) - 10} more events (see full report)")

        # ===== Footer =====
        doc.add_page_break()

        doc.add_heading("Disclaimer", 1)
        disclaimer = """This report documents a controlled security assessment conducted in a lab environment with explicit authorization. All activities were logged and auditable. This report is for educational and defensive security purposes only."""
        doc.add_paragraph(disclaimer)

        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f"Report generated by PurpleForge v{report_data.get('tool_version', '0.1.0')}").italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        logger.info(f"DOCX report exported to {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        return None


def check_docx_support() -> bool:
    """Check if DOCX export is available."""
    return _has_docx()
